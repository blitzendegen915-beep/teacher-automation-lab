from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── 費用計算（コードで算出。手計算値を直書きしない）──────────
# 基礎データ.md記載のレート（既に保護者へ配布済みの合宿案内書と同一根拠）を使用。
# グラウンド使用料2,360円・帰路バスのみ5,180円は基礎データ.mdの確定値をそのまま用いる
# （部の会計計算Excelには異なる分母のグラウンド使用料2,310円・バス代5,170円が
# 見られたが、そちらは42名確定前の古い分母(43名)が残っている可能性があるため、
# 既に配布済み・ファクトチェック済みの基礎データ.mdの値を正とする）。

LODGING_UNIT = 9170     # 1泊3食
EXTRA_LUNCH = 1100      # 増昼食（対象日: 8/7）
BBQ = 1100              # BBQ（対象日: 8/6）
BUS_RETURN_ONLY = 5180  # 帰路バスのみ（基礎データ.md記載の個別精算基準）
MISC_PER_DAY = 1000     # 雑費（参加日数で日割り）
GROUND = 2360           # グラウンド使用料（一律）
RESERVE = 1000          # 予備費（一律）
COACH_FEE = 1030        # コーチ費（選手のみ）


def calc_total(nights, days_count, is_player, bbq=True, extra_lunch=True):
    lodging = LODGING_UNIT * nights
    if bbq:
        lodging += BBQ
    if extra_lunch:
        lodging += EXTRA_LUNCH
    coach = COACH_FEE if is_player else 0
    total = lodging + BUS_RETURN_ONLY + MISC_PER_DAY * days_count + GROUND + RESERVE + coach
    return dict(lodging=lodging, bus=BUS_RETURN_ONLY, misc=MISC_PER_DAY * days_count,
                ground=GROUND, reserve=RESERVE, coach=coach, total=total)


STUDENTS = [
    dict(name='吉田　青空', klass='1年E組', kubun='選手',
         join_date='8月5日（水）夕食〜', last_date='8月7日（金）',
         nights=2, days_count=3,
         calc=calc_total(nights=2, days_count=3, is_player=True)),
    dict(name='椎名　薫', klass='2年F組', kubun='マネージャー',
         join_date='8月4日（火）夕食〜', last_date='8月7日（金）',
         nights=3, days_count=4,
         calc=calc_total(nights=3, days_count=4, is_player=False)),
]

assert STUDENTS[0]['calc']['total'] == 33110
assert STUDENTS[1]['calc']['total'] == 42250

# ── 文書生成 ──────────────────────────────────────────────
doc = Document()

section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.8)

doc.styles['Normal'].font.name = '游明朝'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')


def para(text='', align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, size=11,
         space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = '游明朝'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')
    return p


def set_cell(table, row_idx, col_idx, text, bold=False):
    cell = table.cell(row_idx, col_idx)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    run.font.name = '游明朝'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')


for idx, s in enumerate(STUDENTS):
    c = s['calc']

    para('2026年7月17日', align=WD_ALIGN_PARAGRAPH.RIGHT, size=11, space_after=2)
    para('駒澤大学高等学校　ラグビー部顧問　占部涼也　畠山和真', align=WD_ALIGN_PARAGRAPH.RIGHT, size=10.5)

    para(f'{s["klass"]}　{s["name"]}　様　保護者様', size=11, space_before=4, space_after=12)

    para('夏合宿費のご案内（途中参加者用）', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=10)

    body1 = doc.add_paragraph()
    body1.paragraph_format.space_after = Pt(8)
    body1.paragraph_format.first_line_indent = Pt(11)
    run1 = body1.add_run(
        '平素より、ラグビー部の活動にご理解・ご協力をいただき、誠にありがとうございます。'
        'このたびの夏合宿につきまして、途中からのご参加となりますため、参加日程に応じて'
        '合宿費を個別に算出いたしました。下記のとおりご案内申し上げますので、ご確認のうえ'
        'お振込みくださいますようお願い申し上げます。'
    )
    run1.font.size = Pt(11)
    run1.font.name = '游明朝'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

    para('記', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, space_before=4, space_after=6)

    para('１．参加日程', bold=True, size=11, space_before=4, space_after=2)
    para(f'　　合流：{s["join_date"]}　／　最終日：{s["last_date"]}（{s["nights"]}泊）', size=11, space_after=8)

    para('２．費用内訳', bold=True, size=11, space_before=4, space_after=2)

    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    col_widths = [Cm(7.0), Cm(8.4)]
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            cell.width = col_widths[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)

    set_cell(table, 0, 0, '宿泊・食事代')
    set_cell(table, 0, 1, f'{c["lodging"]:,}円　（{s["nights"]}泊＋増昼食＋BBQ）')
    set_cell(table, 1, 0, '交通費（帰路バスのみ）')
    set_cell(table, 1, 1, f'{c["bus"]:,}円　（行きのバス代は対象外）')
    set_cell(table, 2, 0, '雑費')
    set_cell(table, 2, 1, f'{c["misc"]:,}円　（1,000円×参加{s["days_count"]}日分）')
    set_cell(table, 3, 0, 'グラウンド使用料')
    set_cell(table, 3, 1, f'{c["ground"]:,}円')
    set_cell(table, 4, 0, '予備費')
    set_cell(table, 4, 1, f'{c["reserve"]:,}円　（未使用の場合は返金）')
    set_cell(table, 5, 0, 'コーチ費' + ('（選手のみ）' if s['kubun'] == '選手' else ''))
    set_cell(table, 5, 1, f'{c["coach"]:,}円' if c['coach'] else '対象外（マネージャー）')
    set_cell(table, 6, 0, '合計金額', bold=True)
    set_cell(table, 6, 1, f'{c["total"]:,}円', bold=True)

    para('', space_before=8)

    para('３．お振込みについて', bold=True, size=11, space_before=4, space_after=2)
    body2 = doc.add_paragraph()
    body2.paragraph_format.space_after = Pt(8)
    run2 = body2.add_run(
        '三菱UFJ銀行　世田谷支店（130）　普通0887504　「駒澤大学高校ラグビー部　畠山和真」\n'
        f'振込依頼人名：「{s["klass"][0]}年　{s["name"]}」\n'
        '振込期限：合流日の前日までにお願いいたします。'
    )
    run2.font.size = Pt(11)
    run2.font.name = '游明朝'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

    body3 = doc.add_paragraph()
    body3.paragraph_format.space_after = Pt(8)
    run3 = body3.add_run(
        'ご不明な点がございましたら、下記顧問までご連絡ください。'
    )
    run3.font.size = Pt(11)
    run3.font.name = '游明朝'
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

    para('【お問い合わせ先】', bold=True, size=10.5, space_before=4, space_after=4)
    for line in ['駒澤大学高等学校　ラグビー部顧問　占部涼也', 'Mail：ryoyaurabe@gmail.com', 'Tel（携帯）：080-5032-9150']:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        run.font.size = Pt(10.5)
        run.font.name = '游明朝'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

    para('以　上', align=WD_ALIGN_PARAGRAPH.RIGHT, size=11, space_before=12)

    if idx < len(STUDENTS) - 1:
        doc.add_page_break()

out_path = '/home/user/teacher-automation-lab/output/2026_夏合宿費のご案内_途中参加者用.docx'
os.makedirs('/home/user/teacher-automation-lab/output', exist_ok=True)
doc.save(out_path)
print(f'保存完了: {out_path}')
for s in STUDENTS:
    print(s['name'], s['calc'])
