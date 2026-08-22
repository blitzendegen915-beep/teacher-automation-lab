"""docx帳票の生成。

- write_emergency_sheet   : 緊急時対応シート（指導者携行用・学年別一覧）
- write_allergy_list      : アレルギー一覧（宿泊先厨房向け）
- write_non_responders    : 未回答者一覧（督促用）

すべて要配慮個人情報を含みうる。output/anzen/ に出力し、.gitignore で
バージョン管理から除外している。
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

from . import model

ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "output" / "anzen"

FOOTER_TEXT = "取扱注意／個人情報を含む。合宿終了後は回収し廃棄すること。"

FLAG_SHADE = "F8CBAD"  # 薄いオレンジ（アレルギー・既往症の行）
HEADER_SHADE = "DDEBF7"
EMPTY_BOX_SHADE = "FFFF00"  # 最寄り救急記入欄（黄色網掛け）

DASH = "－"


def _set_page_a4(doc: Document, margin_mm: int = 15) -> None:
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(margin_mm)
        section.bottom_margin = Mm(margin_mm)
        section.left_margin = Mm(margin_mm)
        section.right_margin = Mm(margin_mm)


def _shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False, size: int = 9, color: str = None, blank_ok: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if not text and not blank_ok:
        text = DASH
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ""
        run = p.add_run(FOOTER_TEXT)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("999999")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _title(doc: Document, text: str, size: int = 16) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


# ---------------------------------------------------------------------------
# 1. 緊急時対応シート
# ---------------------------------------------------------------------------

EMERGENCY_COLUMNS = [
    ("氏名", 2.0),
    ("アレルギー", 4.3),
    ("病歴・傷歴", 4.3),
    ("連絡先①", 2.4),
    ("連絡先②", 2.4),
    ("連絡先③", 2.4),
]


def write_emergency_sheet(responses: list, path=None) -> Path:
    path = Path(path) if path else OUTPUT_DIR / "緊急時対応シート.docx"
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _set_page_a4(doc)
    _add_footer(doc)

    style = doc.styles["Normal"]
    style.font.size = Pt(9)

    _title(doc, "緊急時対応シート")
    sub = doc.add_paragraph()
    sub_run = sub.add_run("合宿・遠征帯同用／取扱注意（要配慮個人情報を含む）")
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = RGBColor.from_string("C00000")

    # -- 連絡先ヘッダーブロック ---------------------------------------------
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    info_rows = [
        ("学校代表", "03-3700-6131"),
        ("顧問（占部涼也）", "080-5032-9150"),
        ("宿泊先（菅平ホテル）", "0268-74-2001"),
        ("最寄り救急（現地で記入）", ""),
    ]
    for i, (label, value) in enumerate(info_rows):
        _set_cell_text(info_table.cell(i, 0), label, bold=True, size=10)
        if value:
            _set_cell_text(info_table.cell(i, 1), value, size=10)
        else:
            _set_cell_text(info_table.cell(i, 1), "", size=10, blank_ok=True)
            _shade_cell(info_table.cell(i, 1), EMPTY_BOX_SHADE)
            info_table.rows[i].height = Cm(1.2)
        info_table.cell(i, 0).width = Cm(5.0)
        info_table.cell(i, 1).width = Cm(10.0)

    doc.add_paragraph()
    legend = doc.add_paragraph()
    legend_run = legend.add_run("※ オレンジ網掛け・太字＝アレルギーまたは病歴・傷歴の申告あり。")
    legend_run.font.size = Pt(8)

    grouped = model.group_by_grade(responses)
    for grade, rows in grouped.items():
        doc.add_paragraph()
        gp = doc.add_paragraph()
        grun = gp.add_run(f"{grade}年生（{len(rows)}名）")
        grun.bold = True
        grun.font.size = Pt(12)

        table = doc.add_table(rows=1, cols=len(EMERGENCY_COLUMNS))
        table.style = "Table Grid"
        table.autofit = False
        for c, (label, width) in enumerate(EMERGENCY_COLUMNS):
            cell = table.cell(0, c)
            _set_cell_text(cell, label, bold=True, size=9)
            _shade_cell(cell, HEADER_SHADE)
            cell.width = Cm(width)

        for r in rows:
            row_cells = table.add_row().cells
            flagged = r.is_flagged
            values = [
                r.name,
                ("⚠ " + r.allergy) if r.has_allergy else DASH,
                ("⚠ " + r.medical) if r.has_medical else DASH,
                r.contact1,
                r.contact2,
                r.contact3,
            ]
            for c, val in enumerate(values):
                _set_cell_text(row_cells[c], val, bold=flagged, size=9)
                row_cells[c].width = Cm(EMERGENCY_COLUMNS[c][1])
                if flagged:
                    _shade_cell(row_cells[c], FLAG_SHADE)

    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# 2. アレルギー一覧（宿泊先厨房向け）
# ---------------------------------------------------------------------------


def write_allergy_list(responses: list, path=None) -> Path:
    path = Path(path) if path else OUTPUT_DIR / "アレルギー一覧.docx"
    path.parent.mkdir(parents=True, exist_ok=True)

    allergy_rows = [r for r in responses if r.has_allergy]
    allergy_rows.sort(key=lambda r: (r.grade, r.klass or "", r.sort_number, r.name))

    doc = Document()
    _set_page_a4(doc)
    _add_footer(doc)

    _title(doc, "アレルギー一覧（お食事ご担当者様）")
    sub = doc.add_paragraph()
    sub.add_run(f"該当者：{len(allergy_rows)}名").bold = True

    note = doc.add_paragraph()
    note_run = note.add_run(
        "※ 本一覧は食物アレルギーの申告内容のみを記載しています。既往症・服薬情報・連絡先は含みません。"
    )
    note_run.font.size = Pt(8)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    _set_cell_text(header_cells[0], "氏名", bold=True, size=10)
    _set_cell_text(header_cells[1], "アレルギー内容", bold=True, size=10)
    for c in header_cells:
        _shade_cell(c, HEADER_SHADE)
    header_cells[0].width = Cm(4.0)
    header_cells[1].width = Cm(13.0)

    for r in allergy_rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], r.name, bold=True, size=10)
        _set_cell_text(cells[1], r.allergy, size=10)
        cells[0].width = Cm(4.0)
        cells[1].width = Cm(13.0)

    if not allergy_rows:
        p = doc.add_paragraph()
        p.add_run("該当者はいません。").italic = True

    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# 3. 未回答者一覧
# ---------------------------------------------------------------------------


def write_non_responders(non_responders: list, path=None) -> Path:
    path = Path(path) if path else OUTPUT_DIR / "未回答者一覧.docx"
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _set_page_a4(doc)
    _add_footer(doc)

    _title(doc, "未回答者一覧（安全調査アンケート）")
    sub = doc.add_paragraph()
    sub.add_run(f"未回答：{len(non_responders)}名").bold = True

    if not non_responders:
        p = doc.add_paragraph()
        p.add_run("全員が回答済みです。").italic = True
        doc.save(path)
        return path

    grouped = model.group_roster_by_grade_class(non_responders)
    for (grade, klass), students in grouped.items():
        doc.add_paragraph()
        gp = doc.add_paragraph()
        grun = gp.add_run(f"{grade}年 {klass}組（{len(students)}名）")
        grun.bold = True
        grun.font.size = Pt(11)

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        _set_cell_text(header_cells[0], "氏名", bold=True, size=10)
        _set_cell_text(header_cells[1], "役割", bold=True, size=10)
        for c in header_cells:
            _shade_cell(c, HEADER_SHADE)
        header_cells[0].width = Cm(6.0)
        header_cells[1].width = Cm(4.0)

        for s in students:
            cells = table.add_row().cells
            _set_cell_text(cells[0], s.name, size=10)
            _set_cell_text(cells[1], s.role, size=10)
            cells[0].width = Cm(6.0)
            cells[1].width = Cm(4.0)

    doc.save(path)
    return path
