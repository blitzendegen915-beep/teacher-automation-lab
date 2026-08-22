"""scripts.anzen の単体テスト。pytest不要、`python3 tests/test_anzen.py` で実行できる。

すべて一時ディレクトリに対して実行し、data/rugby/ や output/ には一切書き込まない
（survey-sample.csv はコミット済みのサンプルとして読み込み専用で使う）。
"""
import csv
import shutil
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from scripts.anzen import model as m  # noqa: E402
from scripts.anzen import reports  # noqa: E402

PASS = 0
FAIL = 0


def check(label, condition):
    global PASS, FAIL
    if condition:
        PASS += 1
        print(f"OK   {label}")
    else:
        FAIL += 1
        print(f"FAIL {label}")


def _write_csv(path: Path, headers: list, rows: list) -> None:
    with open(path, "w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(headers)
        for row in rows:
            w.writerow(row)


def test_header_substring_matching(tmpdir: Path):
    """Googleフォームらしい冗長な設問文・列順の入れ替え・紛らわしい語（電話'番号'）でも
    正しい列にマッピングされることを確認する。"""
    headers = [
        "氏名をフルネームで記入してください",
        "緊急連絡先①（電話番号）を記入してください",  # 「番号」を含むが contact1 に割り当たるべき
        "タイムスタンプ",
        "組（アルファベット1文字）を教えてください",
        "病歴・傷歴や既往症、服用中の薬があれば記入してください",
        "学年を選択してください",
        "アレルギーの有無を具体的に記入してください（無ければ「なし」）",
        "出席番号を教えてください",
        "緊急連絡先②（電話番号）を記入してください",
        "緊急連絡先③（電話番号）を記入してください",
    ]
    field_map = m.map_headers(headers)

    check("氏名列が正しくマッピングされる", field_map.get("name") == headers[0])
    check(
        "緊急連絡先①が「番号」フィールドと誤認されずcontact1にマッピングされる",
        field_map.get("contact1") == headers[1],
    )
    check("緊急連絡先②が正しくマッピングされる", field_map.get("contact2") == headers[8])
    check("緊急連絡先③が正しくマッピングされる", field_map.get("contact3") == headers[9])
    check(
        "出席番号（番号）列がcontact系と別にnumberへマッピングされる",
        field_map.get("number") == headers[7],
    )
    check("学年列が正しくマッピングされる", field_map.get("grade") == headers[5])
    check("組列が正しくマッピングされる", field_map.get("klass") == headers[3])
    check("病歴列が正しくマッピングされる", field_map.get("medical") == headers[4])
    check("アレルギー列が正しくマッピングされる", field_map.get("allergy") == headers[6])
    check("タイムスタンプ列が正しくマッピングされる", field_map.get("timestamp") == headers[2])

    # 実際にCSVとして読み込んでも同じ結果になることを確認する。
    path = tmpdir / "survey.csv"
    _write_csv(
        path,
        headers,
        [
            [
                "見本 一郎",
                "090-0000-0001",
                "2026/08/01 09:00:00",
                "A",
                "なし",
                "1",
                "なし",
                "3",
                "",
                "",
            ]
        ],
    )
    result = m.load_survey(path)
    check("CSV読み込み後も1件の回答になる", len(result.responses) == 1)
    r = result.responses[0]
    check("読み込んだ回答のcontact1が電話番号(090-0000-0001)である(numberと混同しない)", r.contact1 == "090-0000-0001")
    check("読み込んだ回答のnumberが'3'である", r.number == "3")


def test_timestamp_parsing():
    check(
        "スラッシュ区切りのタイムスタンプを解析できる",
        m.parse_timestamp("2026/08/01 09:00:00") is not None,
    )
    check(
        "ハイフン区切りのタイムスタンプを解析できる",
        m.parse_timestamp("2026-08-01 09:00:00") is not None,
    )
    check("空文字は解析できずNoneになる", m.parse_timestamp("") is None)
    check("壊れた文字列は解析できずNoneになる", m.parse_timestamp("not-a-date") is None)


def test_duplicate_resolution_sample_csv():
    """コミット済みのサンプルCSV(架空データ)で重複解決を検証する。"""
    result = m.load_survey(m.DEFAULT_SURVEY_PATH)
    check("サンプルCSVは5行ある", result.total_rows == 5)
    check("重複解決後は4名になる(見本太郎が2件→1件)", len(result.responses) == 4)
    check("破棄件数は1件", result.dropped_count == 1)
    check("重複リストに見本太郎が1件ある", len(result.duplicates) == 1 and result.duplicates[0].name == "見本 太郎")
    check("重複リストのcountは2", result.duplicates[0].count == 2)

    taro = next(r for r in result.responses if r.name == "見本 太郎")
    check(
        "重複解決は最新タイムスタンプの内容(連絡先②③が追加された版)を採用する",
        taro.contact2 == "080-2222-3333" and taro.contact3 == "090-9999-0000",
    )

    hanako = next(r for r in result.responses if r.name == "見本 花子")
    check("見本花子はアレルギーありと判定される", hanako.has_allergy is True)
    jiro = next(r for r in result.responses if r.name == "見本 次郎")
    check("見本次郎は病歴ありと判定される", jiro.has_medical is True)
    check("見本次郎はアレルギーなしと判定される('なし')", jiro.has_allergy is False)
    saburo = next(r for r in result.responses if r.name == "見本 三郎")
    check(
        "見本三郎は緊急連絡先①のみ入力・②③は空",
        saburo.contact1 and not saburo.contact2 and not saburo.contact3,
    )


def test_duplicate_resolution_synthetic(tmpdir: Path):
    """タイムスタンプの前後関係が読み込み順と逆でも、最新の内容が採用されることを確認する。"""
    headers = ["タイムスタンプ", "学年", "組", "番号", "氏名", "アレルギー調査", "病歴・傷歴", "緊急連絡先①", "緊急連絡先②", "緊急連絡先③"]
    rows = [
        ["2026/08/05 10:00:00", "2", "A", "1", "見本 逆転", "そばアレルギー", "なし", "090-1111-1111", "", ""],
        ["2026/08/01 08:00:00", "2", "A", "1", "見本 逆転", "なし", "なし", "090-2222-2222", "", ""],
    ]
    path = tmpdir / "reverse.csv"
    _write_csv(path, headers, rows)
    result = m.load_survey(path)
    check("時系列が逆順でCSVに並んでいても重複は1件解決される", len(result.responses) == 1)
    r = result.responses[0]
    check(
        "CSV内の並び順に関わらず、タイムスタンプが新しい方(そばアレルギーあり)が採用される",
        r.has_allergy is True and r.contact1 == "090-1111-1111",
    )


def test_non_responder_detection(tmpdir: Path):
    roster_path = tmpdir / "roster.csv"
    _write_csv(
        roster_path,
        ["grade", "class", "name", "role", "attends", "join_date", "join_meal", "leave_date", "leave_meal"],
        [
            ["1", "A", "見本 回答済", "選手", "yes", "", "", "", ""],
            ["1", "A", "見本 未回答", "選手", "yes", "", "", "", ""],
            ["2", "B", "見本　全角スペース", "選手", "yes", "", "", "", ""],
        ],
    )
    survey_path = tmpdir / "survey.csv"
    _write_csv(
        survey_path,
        ["タイムスタンプ", "学年", "組", "番号", "氏名", "アレルギー調査", "病歴・傷歴", "緊急連絡先①", "緊急連絡先②", "緊急連絡先③"],
        [
            ["2026/08/01 09:00:00", "1", "A", "1", "見本 回答済", "なし", "なし", "090-0000-0000", "", ""],
            # 全角/半角スペースの差異があっても同一人物として名寄せされることを確認する
            ["2026/08/01 09:00:00", "2", "B", "1", "見本 全角スペース", "なし", "なし", "090-0000-0001", "", ""],
        ],
    )

    roster = m.load_roster(roster_path)
    result = m.load_survey(survey_path)
    non_responders = m.find_non_responders(roster, result.responses)

    check("名簿3名中、未回答は1名だけ検出される", len(non_responders) == 1)
    check("未回答者は「見本 未回答」である", non_responders[0].name == "見本 未回答")
    check(
        "全角/半角スペースの違いがあっても回答済みとして名寄せされる(未回答に含まれない)",
        "見本　全角スペース" not in [s.name for s in non_responders]
        and "見本 全角スペース" not in [s.name for s in non_responders],
    )


def test_allergy_list_excludes_phone_and_medical(tmpdir: Path):
    responses = [
        m.Response(
            grade="1", klass="A", number="1", name="見本 アレルギー太郎",
            allergy="えびアレルギーあり", medical="喘息の既往あり・要注意薬あり",
            contact1="090-1234-5678", contact2="03-1111-2222", contact3="",
            timestamp_raw="2026/08/01 09:00:00", row_index=0,
        ),
        m.Response(
            grade="1", klass="A", number="2", name="見本 健康花子",
            allergy="なし", medical="なし",
            contact1="090-9999-8888", contact2="", contact3="",
            timestamp_raw="2026/08/01 09:00:00", row_index=1,
        ),
    ]
    out_path = tmpdir / "アレルギー一覧.docx"
    reports.write_allergy_list(responses, path=out_path)

    from docx import Document

    doc = Document(out_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            full_text += "\n" + "\n".join(c.text for c in row.cells)

    check("アレルギー一覧に該当者(見本アレルギー太郎)の氏名が含まれる", "見本 アレルギー太郎" in full_text)
    check("アレルギー一覧に該当者のアレルギー内容が含まれる", "えびアレルギーあり" in full_text)
    check(
        "アレルギーなしの生徒(見本健康花子)は一覧に含まれない",
        "見本 健康花子" not in full_text,
    )
    check("病歴の内容(喘息)は一覧に含まれない", "喘息" not in full_text)
    check("電話番号は一覧に含まれない", "090-1234-5678" not in full_text and "03-1111-2222" not in full_text)
    check("該当者数の記載が正しい(1名)", "1名" in full_text)


def test_emergency_sheet_flags_allergy_and_medical_rows(tmpdir: Path):
    responses = [
        m.Response(
            grade="1", klass="A", number="1", name="見本 要注意",
            allergy="そばアレルギー", medical="なし",
            contact1="090-0000-0000", contact2="", contact3="",
            timestamp_raw="2026/08/01 09:00:00", row_index=0,
        ),
        m.Response(
            grade="1", klass="A", number="2", name="見本 健康",
            allergy="なし", medical="なし",
            contact1="090-1111-1111", contact2="", contact3="",
            timestamp_raw="2026/08/01 09:00:00", row_index=1,
        ),
    ]
    out_path = tmpdir / "緊急時対応シート.docx"
    reports.write_emergency_sheet(responses, path=out_path)

    from docx import Document

    doc = Document(out_path)
    check("緊急時対応シートが開ける(壊れていない)", len(doc.paragraphs) > 0)

    # 学年別テーブル(先頭のヘッダー情報テーブルの次)を探す
    grade_table = doc.tables[1]
    rows_text = [[c.text for c in row.cells] for row in grade_table.rows]
    flagged_row = next(r for r in rows_text if r[0] == "見本 要注意")
    normal_row = next(r for r in rows_text if r[0] == "見本 健康")
    check("アレルギーありの生徒の行にアレルギー内容が表示される", "そば" in flagged_row[1])
    check("アレルギーなしの生徒の行はダッシュ表示", normal_row[1] == "－")


def test_generate_writes_to_temp_dir_not_output(tmpdir: Path):
    """reports.* にpathを明示すればoutput/以下に書き込まれないことを確認する。"""
    result = m.load_survey(m.DEFAULT_SURVEY_PATH)
    p1 = reports.write_emergency_sheet(result.responses, path=tmpdir / "e.docx")
    p2 = reports.write_allergy_list(result.responses, path=tmpdir / "a.docx")
    p3 = reports.write_non_responders([], path=tmpdir / "n.docx")
    for p in (p1, p2, p3):
        check(f"{p.name} が一時ディレクトリに作成され output/ 配下ではない", str(p).startswith(str(tmpdir)))
        check(f"{p.name} が実在するファイルである", p.exists())


def main():
    tmpdir = Path(tempfile.mkdtemp(prefix="anzen-test-"))
    try:
        test_header_substring_matching(tmpdir)
        test_timestamp_parsing()
        test_duplicate_resolution_sample_csv()
        test_duplicate_resolution_synthetic(tmpdir)
        test_non_responder_detection(tmpdir)
        test_allergy_list_excludes_phone_and_medical(tmpdir)
        test_emergency_sheet_flags_allergy_and_medical_rows(tmpdir)
        test_generate_writes_to_temp_dir_not_output(tmpdir)
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)

    print(f"\n{PASS} passed, {FAIL} failed")
    return 1 if FAIL else 0


if __name__ == "__main__":
    raise SystemExit(main())
